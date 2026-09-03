# -*- coding: utf-8 -*-
"""final_gate.py — 投稿前终检（全量回归断言，~80 项）

对当前 submission_JGG.docx 断言前九轮审计确立的全部事实。
任何一项 FAIL 都不允许投稿。
"""
import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).parent.parent
d = Document(ROOT / 'docs' / 'submission_JGG.docx')
full = '\n'.join(p.text for p in d.paragraphs)
for tb in d.tables:
    for row in tb.rows:
        for c in row.cells:
            full += '\n' + c.text

def _pkg_abstract_ok():
    md = (ROOT / 'docs' / 'manuscript_JGG.md').read_text(encoding='utf-8')
    pkg = (ROOT / 'docs' / 'submission_package.md').read_text(encoding='utf-8')
    m = re.search(r'## Abstract\n\n(.+?)\n\n## Introduction', md, re.S)
    p = re.search(r'## Abstract[^\n]*\n\n(.+?)\n\n---', pkg, re.S)
    if not (m and p):
        return False
    n = lambda s: ' '.join(s.replace('**', '').split())
    return n(m.group(1)) == n(p.group(1))



def _keywords_ok():
    src = (ROOT / 'scripts' / 'build_docx.py').read_text(encoding='utf-8')
    raw = re.search(r'KEYWORDS\s*=\s*\((.+?)\)', src, re.S)
    if not raw:
        return False
    n = lambda s: [w.strip(' .') for w in re.sub(r'\s+', ' ', s).split(';') if w.strip(' .')]
    kw_docx = n(raw.group(1).replace('"', ''))
    m = re.search(r'Keywords:\s*(.+)', (ROOT / 'docs' / 'manuscript_JGG.md').read_text(encoding='utf-8'))
    kw_md = n(m.group(1).split('|')[0]) if m else []
    p2 = re.search(r'\*\*Keywords:\*\*\s*(.+)', (ROOT / 'docs' / 'submission_package.md').read_text(encoding='utf-8'))
    kw_pkg = n(p2.group(1)) if p2 else []
    return kw_docx == kw_md == kw_pkg and len(kw_md) == 7


checks = {
    # ===== 结构 =====
    '表格数=9': len(d.tables) == 9,
    '无中文残留': not re.search('[\u4e00-\u9fff]', full),
    '无游离图标记': not re.search(r'^\(Fig\.', full, re.M),
    '标题正确': 'Multi-vendor evaluation of large language models for ACMG/AMP variant classification with controlled data contamination' in full,
    'Running title': 'Multi-vendor LLM Variant Classification' in full,
    'AI使用声明(含代码辅助)': 'drafting, language editing, and analysis-code development' in full,

    # ===== 摘要 =====
    '摘要-规模': '30,000 evaluations' in full and '15,000 additional evaluations' in full,
    '摘要-范围': '61.8–71.6%' in full and '86–93%' in full,
    '摘要-盲分层': 'no international model outperforms the best domestic models' in full and 'statistically indistinguishable (81.4% vs. 80.2%' in full,
    '盲分层三Tie': 'Gemini 88.0% [85.7–89.9], Claude 87.1% [84.8–89.1], and the domestic leader Qwen 86.8%' in full,
    '专用盲集S4': '81.4% [79.6–83.0]' in full and '64.7% [62.6–66.8]' in full and '25.0%' in full,
    '专用盲集设计': '2,000 newly sampled variants' in full and 'pool: 4,772' in full,
    '摘要-Claude': '97.0% conditional accuracy with 3.9% FP' in full,
    '摘要-AF': 'up to 60.1 pp' in full,

    # ===== 表1（全部经原始数据复算）=====
    'T1-Gemini': '76.5%' in full and '4,538' in full and '91.0%' in full,
    'T1-Qwen': '71.6%' in full and '3,714' in full and '96.4%' in full,
    'T1-Claude': '68.5%' in full and '3,532' in full,
    'T1-Kimi': '67.0%' in full and '3,421' in full and '97.8%' in full,
    'T1-MiMo': '66.1%' in full and '3,876' in full,
    'T1-V4pro': '61.8%' in full and '3,806' in full,
    'T1-GPT': '60.3%' in full and '3,469' in full,
    'T1-chat': '49.4%' in full and '2,504' in full and '98.6%' in full,
    'T1-coder': '49.2%' in full and '2,494' in full and '98.7%' in full,
    'T1-共识行': '64.1%' in full and '98.3%' in full and '2,915' in full and '93.5%' in full,
    'T1-平局528': 'n=528' in full,
    'T1-Wilson脚注': '[70.4, 72.9]' in full and '[65.6, 68.2]' in full,

    # ===== 发现与统计（复算值）=====
    'F1-代际': '+12.4 to +22.4' in full and 'p ≤ 1.6×10⁻⁸⁶' in full,
    'F2-FP': '28.4%' in full and '22.3%' in full and '4.7%' in full,
    'F2-共识FP1.8': '1.8%' in full,
    'F3-投票': '+7.5 pp' in full,
    '表面线索': 'n=1,671' in full and 'n=828' in full and '99.4%' in full and '83.6%' in full and '−15.8 pp' in full,
    '五分类': '32.1%' in full and '11/900' in full and '7.2%' in full,
    'LB→B 53%': '53% of Likely benign' in full,
    'WESI-V4': '0.585' in full and '711 (28.4%)' in full and '716' in full,
    'WESI-Gemini': '0.577' in full and '694 (27.8%)' in full and '712' in full,
    'WESI-MiMo': '0.461' in full and '557 (22.3%)' in full and '573' in full,
    'WESI-coder': '0.026' in full and '32 (1.3%)' in full and '33' in full,
    'WESI-权重2': 'unparseable output = weight 2' in full,
    '确定性-n200谱系': 'Kimi 96.0% > chat 92.5% > Claude 88.0% > Gemini 86.0% > GPT 78.0%' in full,
    '确定性-翻转': 'GPT = 10, Gemini = 15, V4-pro = 2' in full,
    '确定性-交叉验证': 'chat 99/100, coder 100/100, Kimi 97/100' in full,
    '表S1': '40/50 (80.0%)' in full and '44/50 (88.0%)' in full and '38/50 (76.0%)' in full,

    # ===== 表S2/S3/国际 =====
    'S2-弃权': '9.2%' in full and '29.3%' in full and '50.1%' in full,
    'S2-FP列': '27.8%' in full and '3.9%' in full and '17.9%' in full,
    'S3-五行': '79.0%' in full and '72.8%' in full and '69.6%' in full and '42.9%' in full and '43.3%' in full,
    'S3-全900': 'Kimi 74.7% / chat 45.7% / coder 46.0%' in full,
    'S3-Claude502': 'HTTP 502' in full,
    'F6-Fisher': 'Fisher exact p = 0.008' in full,

    # ===== AF/三角验证（复算值）=====
    'AF-良性侧': '8.7% to 68.8%' in full and '+60.1 pp' in full and '45.5% to 81.5%' in full,
    'AF-弃权': '90% to 31%' in full,
    'AF-P侧反转': '76.7%→64.0%' in full and '90.0%→85.3%' in full,
    'AF-九模型范围': '+32.3 to +60.1 pp' in full,
    'AF-P侧弃权升': '23.3%→36.0%' in full,
    '冲突': '+22.4 pp (Kimi)' in full and '+39.1 pp (chat)' in full,
    'MaveDB': '73–93%' in full and '45–55%' in full,
    '校准': '0.95 for Gemini' in full,
    '提示词对称': '71.6-to-65.5%' in full and 'FP 4.7-to-1.0%' in full and '3,186 co-definitive' in full,
    'LB计数': '1,065/5,000' in full and '646/5,000' in full and '591/5,000' in full,
    'LP零': '0/15,000' in full,

    # ===== 方法与声明 =====
    '月度分布': 'Jan 2,097 / Feb 1,672' in full,
    '基因聚类': '2,050 genes' in full and 'NF1 n=83' in full,
    '抽样构成': '2,499 strict Pathogenic + 1 compound' in full,
    '解析失败4行': '4 rows remain unparseable' in full,
    '专家面板日期': 'between 2026-01 and 2026-07' in full,
    'bootstrap': '1.6–2.4× (mean ≈1.9×)' in full,
    'Claude试点5/20': '5 of 20' in full,
    '可复现声明': 'every intermediate file needed to reproduce' in full,
    '级联假设': 'Assuming 3–5 relatives' in full,
    '仓库地址': 'github.com/zksdu/llm-acmg-variant-audit' in full,
    'S4表存在': 'Table S4' in full and 'LastEvaluated ≥ 2026-04' in full,
    '截止披露': 'GPT-5.6-terra ~2026-02-16' in full and 'Gemini Flash model ~2026-03' in full and 'versions 3.5' in full,
    '无GLM残留': 'GLM/MiMo' not in full,
    '投稿包摘要=正文摘要': _pkg_abstract_ok(),
    '关键词三方一致': _keywords_ok(),
    'Zenodo DOI': '10.5281/zenodo.22281813' in full,

    # ===== 参考文献（全部经权威元数据逐字段核实）=====
    'R-MaveDB': '10.1186/s13059-019-1845-6' in full and 'distribute and interpret' in full and '20, 223' in full,
    'R-ClinGen': '10.1056/NEJMsr1406261' in full,
    'R-mygene': '10.1093/nar/gks1114' in full,
    'R-Lin': '10.1038/s41698-025-00935-4' in full and '9, 141' in full and 'Lin, K.-H.' in full,
    'R-Kimi署名': 'Kimi Team, 2025' in full,
    'R-Qwen署名': 'Yang, A., Li, A., Yang, B., et al., 2025' in full,
    'R-Basharat': 'Basharat, H., Plotkin, S., Le, C., Zhu, K., Pink, M., Alfaro, I., 2025' in full,
    'R-Bordt': 'Pf0PaYS9KG' in full and 'Bordt, S., Srinivas, S., Boreiko, V., von Luxburg, U., 2025' in full,
    'R-AIcura': 'scitranslmed.adz4172' in full,
    'R-引用闭环-模型报告': 'DeepSeek-AI, 2024' in full and 'Kimi: Kimi Team, 2025' in full and 'Yang et al., 2025' in full,
    'R-数据源': 'Esposito et al., 2019' in full and 'Wu et al., 2013' in full,
    'R-Lin正文': 'Lin et al., 2025' in full,

    # ===== 历史错误零残留 =====
    '旧-AF基线': '11.0% to 68.8%' not in full and '44.9%' not in full and '47.5%' not in full,
    '旧-79翻转': 'V4-pro = 79' not in full,
    '旧-80.2/97.4/3.6': 'led internationally (80.2%)' not in full and '97.4%' not in full and '3.6% false' not in full,
    '旧-DOI': '1409004' not in full and 'gks1186' not in full and '1685-3' not in full and 's41669' not in full,
    '旧-共识': '2,911' not in full and '98.4%' not in full,
    '旧-500样本': '500-variant' not in full,
    '旧-五分类0/900': '0/900' not in full,
    '旧-51%': '51% of Likely' not in full,
    '旧-署名': 'Moonshot AI, 2025' not in full and 'Qwen Team, 2025' not in full,
    '旧-日期': '2026-04)' not in full,
    '旧-AM数字': '59.3%' not in full and '27 of 5,000' not in full,
    '旧-1.5倍': '≈1.5×' not in full and '约 1.5' not in full,
    '旧-12.6': '+12.6' not in full,
    '旧-Table2断链': 'Table 2,' not in full and 'set (Table 2)' not in full,
}

fails = [k for k, v in checks.items() if not v]
for k, v in checks.items():
    if not v:
        print(f'✗ FAIL: {k}')
print(f'\n===== 终检结果: {sum(checks.values())}/{len(checks)} 通过 =====')
if fails:
    sys.exit(1)
print('全部断言通过 — 可投稿状态')
