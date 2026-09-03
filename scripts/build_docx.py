# -*- coding: utf-8 -*-
"""
build_docx.py — manuscript_JGG.md → 投稿用 Word（submission_JGG.docx）

处理：标题页、##/### 标题、段落（**粗体**/*斜体*）、Markdown 表格、
引用块跳过、内嵌图片行跳过（图单独上传）。

使用：python build_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS = Path(__file__).parent.parent / "docs"

TITLE = "Multi-vendor evaluation of large language models for ACMG/AMP variant classification with controlled data contamination"
AUTHOR = "Bing Song¹, Kai Zhang²,*"
AFFIL_1 = "¹The Third Affiliated Hospital of Guangzhou Medical University, Guangzhou, Guangdong, China"
AFFIL_2 = "²Guangdong Communication Polytechnic, Guangzhou, Guangdong, China"
EMAIL = "zhangkai@gdcp.edu.cn"
RUNNING = "Multi-vendor LLM Variant Classification"
KEYWORDS = ("variant classification; ACMG/AMP; large language models; data leakage; "
            "ClinVar; reliability audit; temporal blinding")
ABSTRACT = """Background. Large language models (LLMs) are increasingly proposed for ACMG/AMP variant classification, but training corpora include ClinVar and ClinGen, so reported accuracy may reflect label memorization rather than reasoning.

Objective. To audit LLM variant-classification reliability under controlled label leakage, across vendors and evidence conditions.

Methods. On a temporally blinded test set of 5,000 ClinVar variants (all assessed after January 2026), we evaluated six Chinese LLMs (30,000 evaluations) and three international flagships at full scale (15,000 additional evaluations), with independent validation on 900 expert-panel variants.

Results. Current-generation models achieved 61.8\u201371.6% all-inclusive accuracy under temporal blinding, rising to 86\u201393% on expert-panel variants. Conservative models reached 97.8\u201398.7% conditional accuracy with FP rates under 4.7%, while reasoning models reached 81.2\u201385.2% with FP rates up to 28.4%. Providing allele-frequency evidence raised Benign sensitivity by up to 60.1 pp. Gemini 3 Flash led internationally (76.5%); Claude paired 97.0% conditional accuracy with 3.9% FP.

Conclusions. LLM variant interpretation is reliable only under blinded model selection, complete evidence (allele frequency mandatory), and abstention-as-human-review policies."""


INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*)")


def add_runs(par, text):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(tok[1:-1])
            r.italic = True
        else:
            par.add_run(tok)


def add_table(doc, rows):
    cells = [ [c.strip() for c in r.strip().strip("|").split("|")] for r in rows ]
    cells = [r for r in cells if not all(set(c) <= set("-: ") for c in r)]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j in range(ncol):
            txt = row[j] if j < len(row) else ""
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, txt)
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.bold = True


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    for lvl, sz in (("Heading 1", 13), ("Heading 2", 12)):
        st = doc.styles[lvl]
        st.font.name = "Times New Roman"
        st.font.size = Pt(sz)
        st.font.bold = True
        st.font.color.rgb = None

    # ===== 标题页 =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(14)

    for line, bold in [(RUNNING + "  (running title)", False),
                       ("", False), (AUTHOR, True), (AFFIL_1, False),
                       (AFFIL_2, False),
                       ("*Corresponding author. E-mail: " + EMAIL, False),
                       ("", False)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line.startswith("*") and line.endswith("*") and len(line) > 2:
            r = p.add_run(line[1:-1]); r.italic = True
        else:
            add_runs(p, line)

    doc.add_heading("Abstract", level=1)
    for para in ABSTRACT.split("\n\n"):
        p = doc.add_paragraph()
        add_runs(p, para)

    p = doc.add_paragraph()
    r = p.add_run("Keywords: "); r.bold = True
    p.add_run(KEYWORDS)

    doc.add_page_break()

    # ===== 正文 =====
    src = (DOCS / "manuscript_JGG.md").read_text(encoding="utf-8")
    lines = src.splitlines()
    i = 0
    # 跳过文件头（# 标题行、> 引用块、--- 分隔、Abstract 占位与 Keywords 段）
    while i < len(lines):
        l = lines[i]
        if l.startswith("## Introduction"):
            break
        i += 1

    buf = []
    while i < len(lines):
        l = lines[i]
        if l.startswith("|"):
            buf.append(l)
            i += 1
            continue
        if buf:
            add_table(doc, buf)
            buf = []
        if not l.strip() or l.startswith(">") or l.strip() == "---" \
                or l.startswith("![") or l.startswith("# "):
            i += 1
            continue
        if l.startswith("## "):
            doc.add_heading(l[3:].strip(), level=1)
        elif l.startswith("### "):
            doc.add_heading(l[4:].strip(), level=2)
        else:
            p = doc.add_paragraph()
            add_runs(p, l)
        i += 1
    if buf:
        add_table(doc, buf)

    out = DOCS / "submission_JGG.docx"
    doc.save(out)
    print(f"\u2713 {out}")
    d2 = Document(str(out))
    print(f"  \u6bb5\u843d {len(d2.paragraphs)} | \u8868\u683c {len(d2.tables)}")


if __name__ == "__main__":
    main()
