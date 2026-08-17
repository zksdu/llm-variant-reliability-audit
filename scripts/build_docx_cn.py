# -*- coding: utf-8 -*-
"""
build_docx_cn.py — 中文版 markdown → docx → PDF（Word COM）

输入：docs/manuscript_JGG_中文版.md
输出：docs/manuscript_中文版.docx、docs/manuscript_中文版.pdf

排版：正文宋体 11pt，标题黑体（H1 15pt / H2 13pt），表格 9pt，A4。
使用：python build_docx_cn.py
"""
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOCS = Path(__file__).parent.parent / "docs"
SRC = DOCS / "manuscript_JGG_中文版.md"

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*\n]+?\*)")


def set_cn_font(run, font_cn="宋体", font_en="Times New Roman"):
    run.font.name = font_en
    r = run._element.rPr
    if r is not None:
        rFonts = r.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = r.makeelement(qn("w:rFonts"), {})
            r.append(rFonts)
        rFonts.set(qn("w:eastAsia"), font_cn)


def add_runs(par, text, cn="宋体"):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(tok[1:-1]); r.italic = True
        else:
            r = par.add_run(tok)
        set_cn_font(r, cn)


def add_table(doc, rows):
    cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
    cells = [r for r in cells if not all(set(c) <= set("-: ") for c in r)]
    if not cells:
        return
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = "Table Grid"
    for i, row in enumerate(cells):
        for j in range(ncol):
            txt = row[j] if j < len(row) else ""
            cell = t.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, txt)
            for r in p.runs:
                r.font.size = Pt(9)
                if i == 0:
                    r.bold = True


def main():
    doc = Document()
    # 页面 A4
    for sec in doc.sections:
        sec.page_width, sec.page_height = Cm(21), Cm(29.7)
        sec.left_margin = sec.right_margin = Cm(2.5)

    src = SRC.read_text(encoding="utf-8")
    lines = src.splitlines()
    buf = []
    first_h1 = True
    for line in lines:
        if line.startswith("|"):
            buf.append(line)
            continue
        if buf:
            add_table(doc, buf)
            buf = []
        s = line.strip()
        if not s or s.startswith(">") or s == "---":
            continue
        if s.startswith("# ") and first_h1:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(s[2:])
            r.bold = True
            r.font.size = Pt(15)
            set_cn_font(r, "黑体")
            first_h1 = False
        elif s.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(s[2:])
            r.bold = True
            r.font.size = Pt(14)
            set_cn_font(r, "黑体")
        elif s.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(s[3:])
            r.bold = True
            r.font.size = Pt(13)
            set_cn_font(r, "黑体")
        elif s.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(s[4:])
            r.bold = True
            r.font.size = Pt(12)
            set_cn_font(r, "黑体")
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.3
            add_runs(p, s)
    if buf:
        add_table(doc, buf)

    # ===== 尾部附录：文章图片（中文图注）=====
    FIGS = DOCS / "figures_v2"
    captions = [
        ("fig1.png", "图 1. 时间盲法测试集上的多模型表现。A：六个国内模型双口径准确率（n=5,000/模型）；B：同一 500 变异子集上九模型对比（灰色为国际模型）；C：良性→致病假阳性率（对数轴），虚线为六模型共识 1.8%。"),
        ("fig2.png", "图 2. 证据可得性支配可靠性。A：AF 消融（良性富集子集 n=400×3）的良性敏感度；B：致病富集子集（n=150×2）；C：跨证据情境的弃权率（有无 AF、主集 vs 冲突变异、MaveDB 功能任务）。"),
        ("fig3.png", "图 3. 输出确定性与 \"Likely\" 档坍缩。A：temperature 0 下重跑一致率（浅色为国际模型）；B：金标准可能致病/可能良性变异的输出分布（Kimi，专家集）——强度信息极化为全致病或全良性。"),
        ("fig4.png", "图 4. 金标准良性变异的归宿。九模型对 2,500 个良性变异的三段堆积：正确判 Benign（蓝）/ VUS 弃权（灰）/ 误判 Pathogenic（品红）。"),
        ("fig5.png", "图 5. 九模型行为仪表盘。六审计维度色阶 0-100，综合视图支持模型选择。"),
    ]
    doc.add_page_break()
    p = doc.add_paragraph()
    r = p.add_run("附录：文章图片")
    r.bold = True; r.font.size = Pt(14); set_cn_font(r, "黑体")
    for fname, cap in captions:
        doc.add_page_break()
        doc.add_picture(str(FIGS / fname), width=Cm(16))
        p = doc.add_paragraph()
        add_runs(p, cap)
        for r_ in p.runs:
            r_.font.size = Pt(9)

    out_docx = DOCS / "manuscript_中文版.docx"
    doc.save(out_docx)
    print(f"✓ docx: {out_docx}")

    # Word COM → PDF
    import pythoncom
    import win32com.client
    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        d = word.Documents.Open(str(out_docx), ReadOnly=True)
        out_pdf = DOCS / "manuscript_中文版.pdf"
        d.SaveAs2(str(out_pdf), FileFormat=17)  # 17 = wdFormatPDF
        d.Close(False)
        print(f"✓ pdf:  {out_pdf}")
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


if __name__ == "__main__":
    main()
